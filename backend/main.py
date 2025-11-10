import warnings
# Silenciar warnings de Pydantic sobre validate_default antes de cualquier import que pueda dispararlos
warnings.filterwarnings("ignore", message=r".*validate_default.*")
from pydantic.warnings import UnsupportedFieldAttributeWarning
warnings.filterwarnings("ignore", category=UnsupportedFieldAttributeWarning)
warnings.filterwarnings("ignore", category=UnsupportedFieldAttributeWarning, module=r"pydantic.*")
warnings.filterwarnings("ignore", category=UnsupportedFieldAttributeWarning, module=r"llama_index.*")

from fastapi import FastAPI, HTTPException
from fastapi import UploadFile, File, Query
from fastapi import WebSocket, WebSocketDisconnect
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from llama_index.core import VectorStoreIndex, Settings
from llama_index.core.readers import SimpleDirectoryReader
from pdf_processor import PDFProcessor
from llama_index.llms.azure_openai import AzureOpenAI
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.core.llms import ChatMessage as LlamaIndexChatMessage
import os
from dotenv import load_dotenv
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import asyncio
import uuid
from datetime import datetime
import json
import io
import requests
from fastapi.responses import StreamingResponse, JSONResponse
import azure.cognitiveservices.speech as speechsdk
# Video frames & OCR
try:
    import cv2  # opencv-python-headless
    import numpy as np
    from PIL import Image
    import pytesseract
except Exception:
    cv2 = None
    np = None
    Image = None
    pytesseract = None

# Cargar .env desde la raíz del proyecto
ROOT_ENV = Path(__file__).resolve().parents[1] / ".env"
load_dotenv(dotenv_path=ROOT_ENV)

# Rutas y metadatos del documento actual (puede ser PDF/MP4/DOCX/TXT)
INDEX_INPUT_FILES = []  # lista de paths a indexar
CURRENT_DOC_PATH = None
CURRENT_DOC_TYPE = None  # 'pdf' | 'mp4' | 'docx' | 'txt'

app = FastAPI()

# Asegurar directorio 'static' existe antes de montarlo
try:
    os.makedirs("static", exist_ok=True)
except Exception:
    pass

# Configuración CORS para permitir solicitudes desde el frontend
from fastapi.middleware.cors import CORSMiddleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción, especificar orígenes exactos
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Montar directorio estático para servir imágenes
app.mount("/static", StaticFiles(directory="static"), name="static")

# --- Configuración Tesseract si está embebido en backend/Tesseract-OCR ---
try:
    if pytesseract is not None:
        tesseract_dir = Path(__file__).resolve().parent / "Tesseract-OCR"
        tesseract_exe = tesseract_dir / "tesseract.exe"
        if tesseract_exe.exists():
            pytesseract.pytesseract.tesseract_cmd = str(tesseract_exe)
except Exception:
    pass

# --- Extracción de frames + OCR ---
def extract_frames_and_ocr(
    video_path: str,
    out_dir: Path,
    frame_every_secs: int = 3,
    max_frames: int = 200,
    enable_scene_detect: bool = True,
    ocr_lang: str = "spa+eng",
    ocr_engine: str = "azure",
) -> Tuple[List[Path], str]:
    """Extrae frames representativos y realiza OCR. Devuelve (frame_paths, ocr_text).
    Si no hay dependencias disponibles, devuelve listas vacías y texto vacío.
    """
    frames = []
    ocr_text_parts = []
    try:
        if cv2 is None or Image is None:
            return frames, ""
        out_dir.mkdir(parents=True, exist_ok=True)
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return frames, ""
        fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        total_secs = (total_frames / fps) if fps > 0 else 0

        # Determinar índices de frames según reglas solicitadas
        indices = []
        if max_frames and max_frames > 0:
            # Deseados por intervalo
            try:
                import math
                desired_by_interval = max(1, int(math.ceil(total_secs / max(1, frame_every_secs))))
            except Exception:
                desired_by_interval = max(1, int((total_secs // max(1, frame_every_secs)) + 1))
            desired = min(max_frames, desired_by_interval)
            # Distribuir uniformemente a lo largo de todo el video
            for i in range(desired):
                pos = int(min(total_frames - 1, round((i * (total_frames - 1)) / max(1, desired - 1)))) if desired > 1 else 0
                indices.append(pos)
        else:
            # Sin límite: todos los frames cada FRAME_EVERY_SECS
            step = max(int(fps * max(1, frame_every_secs)), 1)
            indices = list(range(0, total_frames, step))

        # Umbral para detectar cambios de escena (simple, por histograma)
        prev_hist = None
        grabbed = 0
        for idx in indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ok, frame = cap.read()
            if not ok:
                continue
            take = True
            if enable_scene_detect:
                try:
                    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                    hist = cv2.calcHist([gray], [0], None, [64], [0, 256])
                    hist = cv2.normalize(hist, hist).flatten()
                    if prev_hist is not None:
                        diff = float(cv2.compareHist(prev_hist, hist, cv2.HISTCMP_BHATTACHARYYA))
                        take = diff > 0.25
                    prev_hist = hist
                except Exception:
                    pass
            if not take:
                continue
            # Guardar frame
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil = Image.fromarray(rgb)
            frame_name = f"frame_{grabbed:05d}.jpg"
            frame_path = out_dir / frame_name
            pil.save(frame_path, format="JPEG", quality=85)
            frames.append(frame_path)
            # OCR según motor seleccionado
            text_out = None
            try:
                if (ocr_engine or '').lower() == 'tesseract':
                    if pytesseract is not None:
                        print(f"[OCR] Procesando frame {grabbed} con Tesseract...")
                        text_out = pytesseract.image_to_string(pil, lang=ocr_lang)
                else:
                    print(f"[OCR] Procesando frame {grabbed} con Azure OCR...")
                    text_out = azure_ocr_with_gpt4o(pil)
            except Exception as e:
                print(f"[OCR] Error en frame {grabbed}: {str(e)}")
                # Fallback a Tesseract si Azure falla
                try:
                    if pytesseract is not None:
                        print(f"[OCR] Intentando fallback a Tesseract para frame {grabbed}...")
                        text_out = pytesseract.image_to_string(pil, lang=ocr_lang)
                except Exception as e2:
                    print(f"[OCR] Fallback también falló: {str(e2)}")
                    text_out = None
            text_out = (text_out or "").strip()
            if text_out:
                print(f"[OCR] Frame {grabbed}: Texto extraído ({len(text_out)} caracteres)")
                ocr_text_parts.append(text_out)
            else:
                print(f"[OCR] Frame {grabbed}: Sin texto detectado")
            grabbed += 1
        cap.release()
    except Exception:
        pass
    return frames, "\n".join(ocr_text_parts).strip()

def azure_ocr_with_gpt4o(pil_img: "Image.Image") -> str:
    try:
        import base64, json, requests
        buf = io.BytesIO()
        pil_img.save(buf, format="JPEG", quality=90)
        b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
        api_key = os.getenv("AZURE_OPENAI_API_KEY") or os.getenv("AZURE_FOUNDRY_API_KEY")
        if not endpoint or not api_key:
            print("[Azure OCR] ERROR: Endpoint o API key no configurados")
            raise RuntimeError("Azure OpenAI endpoint/key not configured")
        deployment = os.getenv("AZURE_OCR_DEPLOYMENT", os.getenv("AZURE_OPENAI_GPT4V_DEPLOYMENT", "gpt-4o-mini"))
        api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
        url = f"{endpoint}/openai/deployments/{deployment}/chat/completions?api-version={api_version}"
        print(f"[Azure OCR] Llamando a: {deployment}")
        headers = {"Content-Type": "application/json", "api-key": api_key}
        system_prompt = """Eres un sistema OCR preciso. Tu única tarea es extraer TODO el texto visible en la imagen.

INSTRUCCIONES:
- Extrae ÚNICAMENTE el texto legible (palabras, números, símbolos de texto)
- Respeta el idioma original del texto
- Mantén el formato y estructura del texto lo mejor posible
- Ignora completamente: personas, caras, objetos, imágenes de fondo, elementos visuales no textuales
- NO agregues explicaciones, comentarios o descripciones
- NO menciones qué ves en la imagen (personas, caras, escenarios, etc)
- Si no hay texto visible, devuelve una cadena vacía
- Devuelve SOLAMENTE el texto extraído, nada más"""
        
        payload = {
            "messages": [
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                    ],
                },
            ],
            "temperature": 0.0,
            "max_tokens": 1024,
        }
        resp = requests.post(url, headers=headers, data=json.dumps(payload), timeout=45)
        resp.raise_for_status()
        data = resp.json()
        # Azure responses shape for chat.completions
        result = (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        print(f"[Azure OCR] Respuesta recibida: {len(result)} caracteres")
        return result
    except Exception as e:
        print(f"[Azure OCR] ERROR: {str(e)}")
        import traceback
        print(f"[Azure OCR] Traceback: {traceback.format_exc()}")
        return ""

class ChatMessage(BaseModel):
    role: str  # "user" o "assistant"
    content: str
    timestamp: str = Field(default_factory=lambda: datetime.now().isoformat())
    intent: Optional[str] = None  # depuración: 'rag' o 'generate'

class SourceNode(BaseModel):
    text: str
    page_number: int
    page_image: str = ""  # URL de la imagen de la página
    metadata: Dict[str, Any] = {}

class ChatRequest(BaseModel):
    session_id: Optional[str] = None
    message: str
    provider: str = "azure"  # Por defecto Azure/Foundry

class ChatResponse(BaseModel):
    session_id: str
    message: ChatMessage
    sources: List[SourceNode] = []
    history: List[ChatMessage] = []

class TTSRequest(BaseModel):
    text: str
    voice: Optional[str] = None

# Diccionario para almacenar las sesiones de chat
chat_sessions = {}

# Variables globales para el índice y el proveedor actual
index = None
current_provider = None

def get_default_provider() -> str:
    """Determina proveedor: forzar Azure/Foundry y validar credenciales."""
    key = os.getenv("AZURE_FOUNDRY_API_KEY") or os.getenv("AZURE_OPENAI_API_KEY")
    if not key:
        raise RuntimeError(
            "Falta AZURE_FOUNDRY_API_KEY (o AZURE_OPENAI_API_KEY) en el .env del proyecto."
        )
    return "azure"

# Inicializar el procesador de PDF (solo cuando el archivo es PDF)
pdf_processor = None
pdf_pages = {}
LAST_DOC_TITLE = "tu documento"

# Carga e indexa el/los archivos con el modelo especificado
def build_index(provider: str):
    print(f"Indexando documento(s) con embeddings de: {provider}")
    if not INDEX_INPUT_FILES:
        raise HTTPException(status_code=400, detail="No hay archivo cargado para indexar.")
    for p in INDEX_INPUT_FILES:
        if not os.path.exists(p):
            raise HTTPException(status_code=400, detail=f"Archivo no encontrado para indexar: {p}")
    documents = SimpleDirectoryReader(input_files=INDEX_INPUT_FILES).load_data()
    # Azure/Foundry: usar AzureOpenAIEmbedding con endpoint y api_key
    azure_endpoint = os.getenv("AZURE_FOUNDRY_ENDPOINT") or os.getenv("AZURE_OPENAI_ENDPOINT")
    azure_api_key = os.getenv("AZURE_FOUNDRY_API_KEY") or os.getenv("AZURE_OPENAI_API_KEY", "")
    azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
    embed_deployment = os.getenv("AZURE_OPENAI_EMBED_DEPLOYMENT", "text-embedding-3-small")
    embed_model_name = os.getenv("AZURE_OPENAI_EMBED_MODEL", embed_deployment)

    if not azure_endpoint or not azure_api_key:
        raise RuntimeError("Faltan AZURE_FOUNDRY_ENDPOINT/AZURE_OPENAI_ENDPOINT o la API key en .env")

    embed_model = AzureOpenAIEmbedding(
        model=embed_model_name,
        deployment_name=embed_deployment,
        azure_endpoint=azure_endpoint,
        api_key=azure_api_key,
        api_version=azure_api_version,
    )
    Settings.embed_model = embed_model

    return VectorStoreIndex.from_documents(documents)

# No indexar al inicio; se indexa cuando el usuario sube un archivo
current_provider = get_default_provider()
index = None
print(f"Backend listo. Proveedor: {current_provider}. Esperando archivo para indexar.")

# --- Utilidad: limpiar título legible a partir del nombre del PDF ---
def _clean_pdf_title_from_name(name: str) -> str:
    base = os.path.splitext(os.path.basename(name))[0]
    # Normalizaciones básicas
    base = base.replace("_", " ")
    # Reemplazar conectores comunes con espacios
    base = base.replace(" - ", " ").replace("-", " ")
    # Quitar contenido entre paréntesis
    import re
    s = re.sub(r"\(.*?\)", " ", base)
    # Eliminar tokens que parezcan códigos (muchos mayúsculas/dígitos) o fechas
    tokens = [t for t in s.split() if t]
    filtered = []
    for t in tokens:
        t_no_punct = re.sub(r"[^\wáéíóúñÁÉÍÓÚÑ]", "", t)
        # descartar si es todo dígitos o tiene 5+ dígitos seguidos (fechas/códigos)
        if re.fullmatch(r"\d+", t_no_punct):
            continue
        if re.search(r"\d{5,}", t_no_punct):
            continue
        # descartar si es upper con dígitos y longitud alta
        if re.fullmatch(r"[A-Z0-9]{6,}", t_no_punct):
            continue
        filtered.append(t)
    # Pasar a frase en minúscula con capitalización inicial
    phrase = " ".join(filtered).strip()
    phrase = phrase.lower()
    # Pequeñas limpiezas finales
    phrase = re.sub(r"\s+", " ", phrase)
    if not phrase:
        phrase = "tu documento"
    return phrase

# --- Derivar título del documento según tipo y contenido/metadata ---
def derive_doc_title(path: str, doc_type: str) -> str:
    try:
        base_title = _clean_pdf_title_from_name(path)
    except Exception:
        base_title = "tu documento"
    title = base_title
    try:
        if doc_type == 'pdf':
            try:
                import fitz  # PyMuPDF
                d = fitz.open(path)
                meta_title = (d.metadata or {}).get('title') or ''
                meta_title = (meta_title or '').strip()
                if meta_title and meta_title.lower() not in ['unknown', 'untitled', '']:  # preferible si es válido
                    title = _clean_pdf_title_from_name(meta_title)
                else:
                    # usar primeras líneas de la primera página
                    first_page = d.load_page(0)
                    text = (first_page.get_text() or '').strip()
                    if text:
                        first_line = text.splitlines()[0].strip()
                        if len(first_line) >= 4:
                            title = _clean_pdf_title_from_name(first_line)
            except Exception:
                pass
        elif doc_type == 'docx':
            try:
                from docx import Document
                doc = Document(path)
                meta_title = (getattr(doc, 'core_properties', None) or {}).title if hasattr(doc, 'core_properties') else None
                if meta_title:
                    meta_title = str(meta_title).strip()
                if meta_title:
                    title = _clean_pdf_title_from_name(meta_title)
                else:
                    for p in doc.paragraphs:
                        t = (p.text or '').strip()
                        if len(t) >= 4:
                            title = _clean_pdf_title_from_name(t)
                            break
            except Exception:
                pass
        elif doc_type == 'txt':
            try:
                with open(path, 'r', encoding='utf-8', errors='ignore') as fh:
                    for line in fh:
                        t = (line or '').strip()
                        if len(t) >= 4:
                            title = _clean_pdf_title_from_name(t)
                            break
            except Exception:
                pass
        # mp4: mantener base_title (derivado de nombre), no hay metadata confiable aquí
    except Exception:
        pass
    return title or "tu documento"

# --- Utilidad: convertir Markdown a texto plano para TTS ---
def markdown_to_plaintext(text: str) -> str:
    import re
    if not text:
        return ""
    s = text
    # Quitar bloques de código con triple backticks
    s = re.sub(r"```[\s\S]*?```", " ", s)
    # Quitar código inline `code`
    s = re.sub(r"`([^`]*)`", r"\1", s)
    # Imágenes ![alt](url) -> alt
    s = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", s)
    # Enlaces [texto](url) -> texto
    s = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", s)
    # Encabezados: quitar prefijos # y espacios
    s = re.sub(r"^\s{0,3}#{1,6}\s+", "", s, flags=re.MULTILINE)
    # Citas: quitar > inicial
    s = re.sub(r"^\s*>\s?", "", s, flags=re.MULTILINE)
    # Listas: quitar bullets/num
    s = re.sub(r"^\s*([-*+]\s+)", "", s, flags=re.MULTILINE)
    s = re.sub(r"^\s*\d+\.[)\s]+", "", s, flags=re.MULTILINE)
    # Tablas: quitar bordes y pipes
    s = re.sub(r"^\s*\|", "", s, flags=re.MULTILINE)
    s = re.sub(r"\|\s*", " ", s)
    s = re.sub(r"^\s*:-{2,}:?\s*$", "", s, flags=re.MULTILINE)
    # Negrita/cursiva/strike: eliminar marcadores **, __, *, _, ~~
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"_([^_]+)_", r"\1", s)
    s = re.sub(r"~~([^~]+)~~", r"\1", s)
    # Regla horizontal
    s = re.sub(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", "", s, flags=re.MULTILINE)
    # Colapsar espacios y líneas múltiples
    s = re.sub(r"\s+", " ", s).strip()
    return s

# --- Utilidades para clasificación de intención (RAG vs generación) ---
def _format_history_for_prompt(history: List[ChatMessage], max_chars: int = 2000) -> str:
    parts: List[str] = []
    # Tomar últimas ~10 entradas para contexto, respetando max_chars
    for msg in history[-10:]:
        role = 'Usuario' if msg.role == 'user' else 'Asistente'
        parts.append(f"{role}: {msg.content}")
    s = "\n".join(parts)
    if len(s) > max_chars:
        s = s[-max_chars:]
    return s

def _classify_intent(llm: AzureOpenAI, history: List[ChatMessage], current_user_question: str) -> str:
    hist = _format_history_for_prompt(history)
    doc_type = str(globals().get('CURRENT_DOC_TYPE') or '')
    doc_title = str(globals().get('LAST_DOC_TITLE') or '')
    prompt = (
        "Eres un clasificador. Devuelve SOLO JSON válido con una clave 'intent' cuyo valor sea 'rag' o 'generate'.\n"
        "Usa 'rag' cuando la respuesta requiera consultar el contenido del documento/video indexado.\n"
        "Usa 'generate' cuando se trate de reformular, reescribir o generar sin necesidad de buscar en el índice.\n"
        "Si hay duda o ambigüedad, devuelve 'rag'.\n\n"
        f"doc_type: {doc_type}\n"
        f"doc_title: {doc_title}\n"
        f"history: {hist}\n"
        f"query: {current_user_question}\n\n"
        "Responde estrictamente este JSON:\n"
        "{\"intent\":\"rag\"} o {\"intent\":\"generate\"}"
    )
    try:
        resp = llm.complete(prompt)
        raw = (getattr(resp, 'text', None) or str(resp) or '').strip()
        if raw.startswith('```'):
            raw = raw.strip('`')
            if raw.startswith('json'):
                raw = raw[4:]
        obj = json.loads(raw)
        val = str(obj.get('intent', '')).strip().lower()
        if val in ('rag', 'generate'):
            print(f"[intent] doc_type={doc_type} decided={val} query={current_user_question[:80]}")
            return val
    except Exception:
        pass

    import unicodedata
    def _norm(s: str) -> str:
        s = (s or '').lower()
        return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

    lower_q = _norm(current_user_question)
    video_rag_kws = [
        "video", "charla", "ponencia", "conferencia", "seminario",
        "manual", "archivo", "libro", "texto",
        "de que trata", "de que va", "de que se trata", "tema del video", "contenido del video",
        "resumen del video"
    ]
    for kw in video_rag_kws:
        if kw in lower_q:
            print(f"[intent-fallback] rag by kw '{kw}' doc_type={doc_type}")
            return 'rag'

    for kw in ["resume", "resumen", "sintet", "reformula", "explica", "items", "vinetas", "bullet", "lista", "parafrasea"]:
        if kw in lower_q:
            if (doc_type == 'mp4') and ("documento" not in lower_q):
                print("[intent-fallback] rag by mp4-summary")
                return 'rag'
            print("[intent-fallback] generate by summary kw")
            return 'generate'

    if doc_type == 'mp4':
        if any(p in lower_q for p in ["de que trata", "de que va", "de que se trata", "tema del video", "contenido del video", "resumen del video"]):
            print("[intent-fallback] rag by mp4 phrase")
            return 'rag'

    print("[intent-default] rag")
    return 'rag'

# === Voz: STT (Whisper en Azure OpenAI) y TTS (Azure Speech) ===
@app.post("/stt")
def speech_to_text(audio: UploadFile = File(...)):
    """Transcribe audio usando Whisper en Azure OpenAI.
    Requiere en .env: AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY, AZURE_OPENAI_API_VERSION, AZURE_OPENAI_WHISPER_DEPLOYMENT
    Acepta formatos comunes: webm/ogg/mp3/mp4/wav
    """
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT") or os.getenv("AZURE_FOUNDRY_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY") or os.getenv("AZURE_FOUNDRY_API_KEY")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
    deployment = os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT", "whisper-1")

    if not endpoint or not api_key or not deployment:
        raise HTTPException(status_code=500, detail="Faltan variables para Whisper en .env")

    url = f"{endpoint}/openai/deployments/{deployment}/audio/transcriptions?api-version={api_version}"

    try:
        # Sugerencias de vocabulario/contexto para mejorar precisión
        hint_terms = [
            'charla', 'conferencia', 'ponencia', 'seminario', 'charla',
            'resumen', 'contenido', 'documento', 'pregunta', 'respuesta',
            'RAG', 'embeddings', 'Azure', 'OpenAI'
        ]
        try:
            if (LAST_DOC_TITLE or '').strip() and (LAST_DOC_TITLE or '').strip().lower() != 'tu documento':
                hint_terms.append(LAST_DOC_TITLE)
        except Exception:
            pass
        prompt_txt = " ".join(hint_terms)
        data = {
            # 'model' no es necesario en Azure (va en la ruta como deployment)
            'response_format': 'json',
            'language': 'es',
            'prompt': prompt_txt
        }
        file_bytes = audio.file.read()
        files = {
            'file': (audio.filename or 'audio.webm', file_bytes, audio.content_type or 'application/octet-stream')
        }
        headers = {
            'api-key': api_key,
        }
        resp = requests.post(url, headers=headers, data=data, files=files, timeout=60)
        if resp.status_code != 200:
            return JSONResponse(status_code=resp.status_code, content={"error": resp.text})
        payload = resp.json()
        text = payload.get('text', '')
        return {"text": text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/save_config")
async def save_config(config: Dict[str, Any]):
    """Guarda configuración en el archivo .env"""
    try:
        env_path = Path(__file__).resolve().parents[1] / ".env"
        if not env_path.exists():
            raise HTTPException(status_code=404, detail="Archivo .env no encontrado")
        
        # Leer contenido actual del .env
        with open(env_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Actualizar valores
        config_map = {
            'FRAME_EVERY_SECS': str(config.get('FRAME_EVERY_SECS', 3)),
            'MAX_FRAMES': str(config.get('MAX_FRAMES', 20)),
            'SCENE_DETECTION': 'true' if config.get('SCENE_DETECTION', False) else 'false',
            'DO_OCR_FRAMES': 'true' if config.get('DO_OCR_FRAMES', False) else 'false',
            'SHOW_SOURCES': 'true' if config.get('SHOW_SOURCES', False) else 'false',
            'SPEAK_REPLY': 'true' if config.get('SPEAK_REPLY', False) else 'false',
            'OCR_ENGINE': str(config.get('OCR_ENGINE', 'azure') or 'azure'),
        }
        
        new_lines = []
        updated_keys = set()
        
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and not line_stripped.startswith('#'):
                for key, value in config_map.items():
                    if line_stripped.startswith(f"{key}="):
                        new_lines.append(f"{key}={value}\n")
                        updated_keys.add(key)
                        break
                else:
                    new_lines.append(line)
            else:
                new_lines.append(line)
        
        # Agregar keys que no existían
        for key, value in config_map.items():
            if key not in updated_keys:
                new_lines.append(f"{key}={value}\n")
        
        # Escribir de vuelta
        with open(env_path, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
        # Recargar variables de entorno en el proceso
        try:
            load_dotenv(dotenv_path=env_path, override=True)
        except Exception:
            pass
        
        print(f"[CONFIG] Configuración guardada en .env: {config_map}")
        return {"message": "Configuración guardada correctamente", "config": config_map}
    except Exception as e:
        print(f"[CONFIG] Error al guardar: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error al guardar configuración: {e}")

@app.get("/config")
def get_config():
    """Devuelve configuración efectiva leída del .env"""
    try:
        # Recargar .env para reflejar últimos cambios
        try:
            load_dotenv(dotenv_path=ROOT_ENV, override=True)
        except Exception:
            pass
        def getenv_bool(key: str, default: bool=False) -> bool:
            val = os.getenv(key)
            if val is None:
                return default
            return str(val).lower() in ("1", "true", "yes", "on")

        cfg = {
            'FRAME_EVERY_SECS': int(os.getenv('FRAME_EVERY_SECS', '3') or 3),
            'MAX_FRAMES': int(os.getenv('MAX_FRAMES', '20') or 20),
            'SCENE_DETECTION': getenv_bool('SCENE_DETECTION', True),
            'DO_OCR_FRAMES': getenv_bool('DO_OCR_FRAMES', False),
            'SHOW_SOURCES': getenv_bool('SHOW_SOURCES', False),
            'SPEAK_REPLY': getenv_bool('SPEAK_REPLY', False),
            'OCR_ENGINE': (os.getenv('OCR_ENGINE', 'azure') or 'azure'),
        }
        return cfg
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error leyendo configuración: {e}")

@app.post("/tts")
def text_to_speech(req: TTSRequest):
    """Convierte texto a voz. Usa Azure Speech si hay credenciales; de lo contrario intenta Azure OpenAI (Foundry) Audio TTS."""
    # 1) Camino preferido: Azure Speech
    speech_key = os.getenv("SPEECH_KEY")
    speech_region = os.getenv("SPEECH_REGION")
    default_voice = os.getenv("SPEECH_TTS_VOICE", "es-ES-AlvaroNeural")
    clean_text = markdown_to_plaintext(req.text)
    if speech_key and speech_region:
        speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=speech_region)
        speech_config.speech_synthesis_voice_name = req.voice or default_voice
        speech_config.set_speech_synthesis_output_format(
            speechsdk.SpeechSynthesisOutputFormat.Audio16Khz32KBitRateMonoMp3
        )
        # SDK reciente requiere activar explicitamente el parlante por defecto.
        # Aunque consumimos los bytes resultantes, habilitarlo evita errores en la inicialización.
        audio_config = speechsdk.audio.AudioOutputConfig(use_default_speaker=True)
        synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=audio_config)
        result = synthesizer.speak_text_async(clean_text).get()
        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
            audio_bytes = result.audio_data
            return StreamingResponse(io.BytesIO(audio_bytes), media_type="audio/mpeg")
        elif result.reason == speechsdk.ResultReason.Canceled:
            cancellation_details = result.cancellation_details
            raise HTTPException(status_code=500, detail=f"TTS cancelado: {cancellation_details.reason} - {cancellation_details.error_details}")
        else:
            raise HTTPException(status_code=500, detail="Error desconocido en síntesis de voz (Speech)")

    # 2) Fallback: Azure OpenAI (Foundry) Audio TTS
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT") or os.getenv("AZURE_FOUNDRY_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY") or os.getenv("AZURE_FOUNDRY_API_KEY")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-07-01-preview")
    tts_deployment = os.getenv("AZURE_OPENAI_TTS_DEPLOYMENT")  # p.ej. "tts" o "gpt-4o-mini-tts"
    voice = req.voice or os.getenv("AZURE_OPENAI_TTS_VOICE", "alloy")
    model = os.getenv("AZURE_OPENAI_TTS_MODEL", None)  # opcional, algunos endpoints lo requieren

    if not endpoint or not api_key or not tts_deployment:
        raise HTTPException(status_code=500, detail="Faltan credenciales para TTS (Speech) y tampoco hay AZURE_OPENAI_TTS_DEPLOYMENT en .env")

    url = f"{endpoint}/openai/deployments/{tts_deployment}/audio/speech?api-version={api_version}"
    try:
        headers = {
            "api-key": api_key,
            "Content-Type": "application/json"
        }
        body = {
            "input": clean_text,
            "voice": voice,
            "format": "mp3"
        }
        if model:
            body["model"] = model
        resp = requests.post(url, headers=headers, json=body, timeout=60)
        if resp.status_code != 200:
            # Diagnóstico en consola
            try:
                print("[TTS Foundry] Non-200:", resp.status_code, resp.text)
            except Exception:
                pass
            raise HTTPException(status_code=resp.status_code, detail=resp.text)
        return StreamingResponse(io.BytesIO(resp.content), media_type="audio/mpeg")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Azure OpenAI TTS error: {e}")

@app.post("/upload")
def upload_file(
    file: UploadFile = File(...),
    frame_every_secs: Optional[int] = Query(None),
    max_frames: Optional[int] = Query(None),
    scene_detection: Optional[bool] = Query(None),
    do_ocr_frames: Optional[bool] = Query(None),
    ocr_engine: Optional[str] = Query(None),
):
    """Acepta mp4/pdf/docx/txt, prepara contenido e indexa."""
    global CURRENT_DOC_PATH, CURRENT_DOC_TYPE, pdf_processor, pdf_pages, index, INDEX_INPUT_FILES, LAST_DOC_TITLE

    allowed_types = {
        "application/pdf": "pdf",
        "video/mp4": "mp4",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }
    content_type = file.content_type or ""
    doc_type = allowed_types.get(content_type)
    if doc_type is None:
        # intentar por extensión si content_type vino vacío
        fname = (getattr(file, "filename", "") or "").lower()
        if fname.endswith(".pdf"): doc_type = "pdf"
        elif fname.endswith(".mp4"): doc_type = "mp4"
        elif fname.endswith(".docx"): doc_type = "docx"
        elif fname.endswith(".txt"): doc_type = "txt"
    if doc_type is None:
        raise HTTPException(status_code=400, detail="Tipo de archivo no soportado. Use mp4, pdf, docx o txt.")

    uploads_dir = Path(__file__).resolve().parent / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    # Guardar archivo original
    suffix = {
        "pdf": ".pdf",
        "mp4": ".mp4",
        "docx": ".docx",
        "txt": ".txt",
    }[doc_type]
    orig_target = uploads_dir / f"{uuid.uuid4()}{suffix}"
    with open(orig_target, "wb") as f:
        f.write(file.file.read())

    # Preparar para indexación
    INDEX_INPUT_FILES = []
    pdf_pages = {}
    pdf_processor = None

    # Inicializar título (se calculará por tipo)
    LAST_DOC_TITLE = "tu documento"

    if doc_type == "pdf":
        CURRENT_DOC_PATH = str(orig_target)
        CURRENT_DOC_TYPE = "pdf"
        INDEX_INPUT_FILES = [CURRENT_DOC_PATH]
        # Generar imágenes de páginas para fuentes
        try:
            pdf_processor = PDFProcessor(CURRENT_DOC_PATH)
            pdf_pages = pdf_processor.extract_page_images()
        except Exception:
            pdf_pages = {}
        # Título desde metadata/contenido
        try:
            LAST_DOC_TITLE = derive_doc_title(CURRENT_DOC_PATH, CURRENT_DOC_TYPE)
        except Exception:
            pass
    elif doc_type == "docx":
        CURRENT_DOC_PATH = str(orig_target)
        CURRENT_DOC_TYPE = "docx"
        # Extraer texto a un .txt temporal para indexar
        try:
            from docx import Document  # python-docx
            doc = Document(CURRENT_DOC_PATH)
            text = []
            for p in doc.paragraphs:
                if p.text:
                    text.append(p.text)
            txt = "\n".join(text).strip()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error procesando DOCX: {e}")
        txt_path = uploads_dir / f"{uuid.uuid4()}.txt"
        txt_path.write_text(txt or "", encoding="utf-8")
        INDEX_INPUT_FILES = [str(txt_path)]
        # Título desde metadata o primeras líneas
        try:
            LAST_DOC_TITLE = derive_doc_title(CURRENT_DOC_PATH, CURRENT_DOC_TYPE)
        except Exception:
            pass
    elif doc_type == "txt":
        CURRENT_DOC_PATH = str(orig_target)
        CURRENT_DOC_TYPE = "txt"
        INDEX_INPUT_FILES = [CURRENT_DOC_PATH]
        try:
            LAST_DOC_TITLE = derive_doc_title(CURRENT_DOC_PATH, CURRENT_DOC_TYPE)
        except Exception:
            pass
    elif doc_type == "mp4":
        CURRENT_DOC_PATH = str(orig_target)
        CURRENT_DOC_TYPE = "mp4"
        # Transcribir con Whisper (Azure OpenAI)
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT") or os.getenv("AZURE_FOUNDRY_ENDPOINT")
        api_key = os.getenv("AZURE_OPENAI_API_KEY") or os.getenv("AZURE_FOUNDRY_API_KEY")
        api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
        deployment = os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT", "whisper-1")
        if not endpoint or not api_key or not deployment:
            raise HTTPException(status_code=500, detail="Faltan variables para Whisper en .env")
        url = f"{endpoint}/openai/deployments/{deployment}/audio/transcriptions?api-version={api_version}"
        try:
            with open(CURRENT_DOC_PATH, "rb") as fh:
                files = {"file": (os.path.basename(CURRENT_DOC_PATH), fh.read(), "video/mp4")}
            data = {"response_format": "json", "language": "es"}
            headers = {"api-key": api_key}
            resp = requests.post(url, headers=headers, data=data, files=files, timeout=120)
            if resp.status_code != 200:
                raise HTTPException(status_code=resp.status_code, detail=f"Whisper error: {resp.text}")
            payload = resp.json()
            transcript = payload.get("text", "")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error transcribiendo MP4: {e}")
        # Guardar transcripción a .txt para indexar
        txt_path = uploads_dir / f"{uuid.uuid4()}.txt"
        txt_path.write_text(transcript or "", encoding="utf-8")
        INDEX_INPUT_FILES = [str(txt_path)]
        # Extraer frames y OCR (si hay dependencias)
        video_uuid = uuid.uuid4().hex[:8]
        frames_dir = Path(__file__).resolve().parent / "static" / "video_frames" / video_uuid
        # Resolver parámetros (query param tiene prioridad, luego env)
        fes = frame_every_secs if frame_every_secs is not None else int(os.getenv("FRAME_EVERY_SECS", "3") or 3)
        mf = max_frames if max_frames is not None else int(os.getenv("MAX_FRAMES", "200") or 200)
        sd = scene_detection if scene_detection is not None else (os.getenv("SCENE_DETECTION", "true").lower() != "false")
        # Determinar si hacer OCR (parámetro tiene prioridad, luego env)
        do_ocr = do_ocr_frames if do_ocr_frames is not None else (os.getenv("DO_OCR_FRAMES", "false").lower() == "true")
        print(f"[VIDEO] Iniciando extracción de frames y OCR...")
        print(f"[VIDEO] Parámetros: frame_every_secs={fes}, max_frames={mf}, scene_detect={sd}, do_ocr={do_ocr}")
        
        frame_paths = []
        ocr_text = ""
        if do_ocr:
            frame_paths, ocr_text = extract_frames_and_ocr(
                CURRENT_DOC_PATH,
                frames_dir,
                frame_every_secs=int(fes),
                max_frames=int(mf),
                enable_scene_detect=bool(sd),
                ocr_lang=os.getenv("OCR_LANG", "spa+eng"),
                ocr_engine=(ocr_engine or os.getenv("OCR_ENGINE", "azure")),
            )
            print(f"[VIDEO] Frames extraídos: {len(frame_paths)}")
            print(f"[VIDEO] Texto OCR extraído: {len(ocr_text)} caracteres")
            ocr_txt_path = None
            if ocr_text:
                ocr_txt_path = uploads_dir / f"{uuid.uuid4()}_ocr.txt"
                ocr_txt_path.write_text(ocr_text, encoding="utf-8")
                INDEX_INPUT_FILES.append(str(ocr_txt_path))
                print(f"[VIDEO] Archivo OCR guardado: {ocr_txt_path.name}")
            else:
                print(f"[VIDEO] ADVERTENCIA: No se extrajo texto OCR del video")
        else:
            print(f"[VIDEO] OCR deshabilitado, solo se usará transcripción de audio")
        # Derivar título desde el nombre original (y no desde la transcripción)
        try:
            LAST_DOC_TITLE = derive_doc_title(CURRENT_DOC_PATH, CURRENT_DOC_TYPE)
        except Exception:
            pass

    # (Re)construir índice
    index = build_index(current_provider)

    # Resetear sesiones de chat
    chat_sessions.clear()

    return {
        "message": "Archivo cargado e indexado correctamente.",
        "doc_path": CURRENT_DOC_PATH,
        "doc_type": CURRENT_DOC_TYPE,
        "doc_title": LAST_DOC_TITLE,
        # compat con frontend anterior
        "pdf_path": CURRENT_DOC_PATH,
        "pdf_title": LAST_DOC_TITLE,
        # imágenes de frames (si se generaron)
        "video_frames": [
            f"/static/video_frames/{frames_dir.name}/{p.name}" for p in (frame_paths if 'frame_paths' in locals() else [])
        ] if CURRENT_DOC_TYPE == "mp4" else [],
    }

@app.post("/chat", response_model=ChatResponse)
def chat(req: ChatRequest):
    global index, current_provider
    
    # Forzar Azure/Foundry como proveedor único
    if current_provider != "azure":
        current_provider = "azure"
    if req.provider != "azure":
        req.provider = "azure"
    # Validar que haya índice construido
    if index is None:
        raise HTTPException(status_code=400, detail="No hay índice construido. Cargue un archivo primero.")
    
    # Configurar el modelo LLM (Azure/Foundry)
    azure_endpoint = os.getenv("AZURE_FOUNDRY_ENDPOINT") or os.getenv("AZURE_OPENAI_ENDPOINT")
    azure_api_key = os.getenv("AZURE_FOUNDRY_API_KEY") or os.getenv("AZURE_OPENAI_API_KEY", "")
    azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
    llm_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")

    llm = AzureOpenAI(
        azure_deployment=llm_deployment,
        azure_endpoint=azure_endpoint,
        api_key=azure_api_key,
        api_version=azure_api_version,
        temperature=0.3,
    )
    
    # Crear o recuperar la sesión de chat
    session_id = req.session_id if req.session_id else str(uuid.uuid4())
    
    if session_id not in chat_sessions:
        # Crear sesión con mensaje de bienvenida genérico fijo
        welcome_text = "¡Bienvenido! Te ayudo a consultar sobre la información disponible. ¿Qué te gustaría saber?"
        welcome_msg = ChatMessage(role="assistant", content=welcome_text)
        welcome_li = LlamaIndexChatMessage(role="assistant", content=welcome_text)
        chat_sessions[session_id] = {
            "messages": [welcome_msg],
            "chat_history": [welcome_li]  # Historial para LlamaIndex
        }
    
    # Añadir el mensaje del usuario al historial
    user_message = ChatMessage(role="user", content=req.message)
    chat_sessions[session_id]["messages"].append(user_message)
    
    # Convertir mensajes al formato de LlamaIndex para el historial de chat
    llama_index_message = LlamaIndexChatMessage(role="user", content=req.message)
    chat_sessions[session_id]["chat_history"].append(llama_index_message)
    
    # Clasificar intención
    try:
        intent = _classify_intent(llm, chat_sessions[session_id]["messages"], req.message)
    except Exception:
        intent = 'rag'

    source_nodes: List[SourceNode] = []
    assistant_text: str = ""

    if intent == 'rag':
        # Responder con RAG (vector search)
        # System prompt con ajuste para videos
        base_sys_prompt = (
            "Eres un asistente que responde basado estrictamente en el contenido indexado. "
            "Si no encuentras la información, dilo claramente. "
        )
        if (globals().get('CURRENT_DOC_TYPE') == 'mp4'):
            base_sys_prompt += (
                "El contenido proviene de la transcripción de un video. Trata 'video' y 'documento' como sinónimos. "
                "No digas frases como 'el documento no menciona un video'. "
                "Cuando sea natural, prefiere referirte al material como 'el video'. "
            )
        chat_engine = index.as_chat_engine(
            chat_mode="context",
            llm=llm,
            similarity_top_k=3,
            chat_history=chat_sessions[session_id]["chat_history"],
            system_prompt=base_sys_prompt
        )
        response = chat_engine.chat(req.message)
        assistant_text = str(response)

        # Extraer fuentes
        nodes_iter = None
        if hasattr(response, 'source_nodes') and response.source_nodes:
            nodes_iter = response.source_nodes
        elif hasattr(response, 'sources') and response.sources:
            nodes_iter = response.sources
        if nodes_iter:
            for item in nodes_iter:
                node = getattr(item, 'node', None) or getattr(item, 'source_node', None) or item
                metadata = getattr(node, 'metadata', {}) or {}
                text_val = getattr(node, 'text', None)
                if text_val is None and hasattr(item, 'raw_output'):
                    text_val = str(item.raw_output)
                if text_val is None:
                    text_val = str(node)
                page_number = 1
                try:
                    if isinstance(metadata, dict):
                        if 'page_label' in metadata:
                            page_number = int(metadata['page_label'])
                        elif 'page' in metadata:
                            page_number = int(metadata['page'])
                except (ValueError, TypeError):
                    pass
                page_image = ""
                if 'pdf_pages' in globals() and isinstance(page_number, int) and page_number in pdf_pages:
                    page_image = pdf_pages[page_number]
                source_nodes.append(
                    SourceNode(
                        text=text_val,
                        page_number=page_number,
                        page_image=page_image,
                        metadata=metadata if isinstance(metadata, dict) else {}
                    )
                )
        # Añadir respuesta a historiales
        assistant_message = ChatMessage(role="assistant", content=assistant_text, intent=intent)
        chat_sessions[session_id]["messages"].append(assistant_message)
        llama_index_response = LlamaIndexChatMessage(role="assistant", content=assistant_text)
        chat_sessions[session_id]["chat_history"].append(llama_index_response)
    else:
        # Generación pura (sin RAG)
        hist_str = _format_history_for_prompt(chat_sessions[session_id]["messages"], max_chars=3000)
        gen_prompt = (
            "Eres un asistente que ayuda con redacción y transformación de texto. "
            "Sigue la instrucción del usuario a continuación. No cites fuentes ni hagas referencia al documento subido.\n\n"
            f"Historial reciente:\n{hist_str}\n\n"
            f"Instrucción del usuario: {req.message}\n\n"
            "Respuesta:"
        )
        gen_resp = llm.complete(gen_prompt)
        assistant_text = (getattr(gen_resp, 'text', None) or str(gen_resp))
        assistant_message = ChatMessage(role="assistant", content=assistant_text, intent=intent)
        chat_sessions[session_id]["messages"].append(assistant_message)
        chat_sessions[session_id]["chat_history"].append(LlamaIndexChatMessage(role="assistant", content=assistant_text))

    return ChatResponse(
        session_id=session_id,
        message=assistant_message,
        sources=source_nodes,
        history=chat_sessions[session_id]["messages"]
    )

@app.get("/sessions/{session_id}", response_model=Dict[str, Any])
def get_session(session_id: str):
    if session_id not in chat_sessions:
        raise HTTPException(status_code=404, detail="Sesión no encontrada")
    
    return {
        "session_id": session_id,
        "messages": chat_sessions[session_id]["messages"]
    }

@app.delete("/sessions/{session_id}")
def delete_session(session_id: str):
    if session_id in chat_sessions:
        del chat_sessions[session_id]
    
    return {"message": "Sesión eliminada correctamente"}

@app.post("/reset")
def reset_state():
    """Resetea el estado del backend como si fuera un inicio fresco de la app."""
    global index, INDEX_INPUT_FILES, CURRENT_DOC_PATH, CURRENT_DOC_TYPE, pdf_processor, pdf_pages, LAST_DOC_TITLE
    # Limpiar índice y documento actual
    index = None
    INDEX_INPUT_FILES = []
    CURRENT_DOC_PATH = None
    CURRENT_DOC_TYPE = None
    pdf_processor = None
    pdf_pages = {}
    # Limpiar sesiones
    chat_sessions.clear()
    # Reiniciar título por defecto
    LAST_DOC_TITLE = "tu documento"
    welcome_text = "¡Bienvenido! Te ayudo a consultar sobre la información disponible. ¿Qué te gustaría saber?"
    return {"message": "Estado reiniciado", "welcome_text": welcome_text}
